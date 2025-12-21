"""Document processing service for DOCX and PDF files."""
import io
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import tiktoken

logger = logging.getLogger(__name__)

# Knowledge categories for novel writing
KNOWLEDGE_CATEGORIES = [
    'draft',
    'concept', 
    'character',
    'chapter',
    'settings',
    'plot',
    'dialogue',
    'research',
    'notes'
]


class DocumentProcessor:
    """Process DOCX and PDF documents."""
    
    def __init__(self):
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from a DOCX file."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from a PDF file."""
        try:
            import pdfplumber
            
            full_text = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
            
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise
    
    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text based on file extension."""
        ext = Path(filename).suffix.lower()
        
        if ext == '.docx':
            return self.extract_text_from_docx(file_content)
        elif ext == '.pdf':
            return self.extract_text_from_pdf(file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        return len(self.encoding.encode(text))
    
    def chunk_text(self, text: str, chunk_size: int = 1000, 
                   overlap: int = 200) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks."""
        tokens = self.encoding.encode(text)
        chunks = []
        
        start = 0
        chunk_index = 0
        
        while start < len(tokens):
            end = min(start + chunk_size, len(tokens))
            chunk_tokens = tokens[start:end]
            chunk_text = self.encoding.decode(chunk_tokens)
            
            chunks.append({
                'index': chunk_index,
                'text': chunk_text,
                'token_count': len(chunk_tokens),
                'start_token': start,
                'end_token': end
            })
            
            chunk_index += 1
            start = end - overlap if end < len(tokens) else end
        
        return chunks
    
    def auto_categorize(self, text: str, filename: str) -> str:
        """Auto-categorize content based on keywords and patterns."""
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        # Check filename first
        for category in KNOWLEDGE_CATEGORIES:
            if category in filename_lower:
                return category
        
        # Check content patterns
        character_keywords = ['character', 'protagonist', 'antagonist', 'appearance', 
                             'personality', 'backstory', '角色', '人物', '主角']
        plot_keywords = ['plot', 'story arc', 'conflict', 'resolution', 'climax',
                        '情節', '故事', '衝突', '結局']
        settings_keywords = ['world', 'setting', 'magic system', 'technology', 'geography',
                            '世界', '設定', '魔法', '科技']
        chapter_keywords = ['chapter', 'scene', 'part', '章節', '場景', '第']
        dialogue_keywords = ['dialogue', 'conversation', 'says', 'asked', 'replied',
                            '對話', '說道', '問道']
        
        # Score each category
        scores = {
            'character': sum(1 for kw in character_keywords if kw in text_lower),
            'plot': sum(1 for kw in plot_keywords if kw in text_lower),
            'settings': sum(1 for kw in settings_keywords if kw in text_lower),
            'chapter': sum(1 for kw in chapter_keywords if kw in text_lower),
            'dialogue': sum(1 for kw in dialogue_keywords if kw in text_lower),
        }
        
        # Check if it looks like a draft (has narrative structure)
        if text.count('\n\n') > 5 and len(text) > 2000:
            scores['draft'] = scores.get('draft', 0) + 3
        
        # Return highest scoring category
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        
        return 'notes'  # Default category


class LongContextManager:
    """Manage long context for LLM queries."""
    
    def __init__(self, max_tokens: int = 32000):
        self.max_tokens = max_tokens
        self.encoding = tiktoken.get_encoding("cl100k_base")
        # Reserve tokens for system prompt and response
        self.reserved_tokens = 4000
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        return len(self.encoding.encode(text))
    
    def build_context(self, 
                      user_query: str,
                      chapters: List[Dict[str, Any]] = None,
                      knowledge: List[Dict[str, Any]] = None,
                      characters: List[Dict[str, Any]] = None,
                      events: List[Dict[str, Any]] = None,
                      conversation_history: List[Dict[str, str]] = None,
                      priority_order: List[str] = None) -> Dict[str, Any]:
        """
        Build context within token limits, prioritizing important information.
        
        Args:
            user_query: The user's question
            chapters: Retrieved chapter content
            knowledge: Retrieved knowledge base entries
            characters: Character information from graph
            events: Timeline events from graph
            conversation_history: Recent conversation history
            priority_order: Order of priority for context types
        
        Returns:
            Dict with optimized context and metadata
        """
        available_tokens = self.max_tokens - self.reserved_tokens
        query_tokens = self.count_tokens(user_query)
        available_tokens -= query_tokens
        
        if priority_order is None:
            priority_order = ['characters', 'events', 'chapters', 'knowledge', 'history']
        
        context = {
            'characters': [],
            'events': [],
            'chapters': [],
            'knowledge': [],
            'history': []
        }
        token_usage = {
            'query': query_tokens,
            'characters': 0,
            'events': 0,
            'chapters': 0,
            'knowledge': 0,
            'history': 0,
            'total': query_tokens
        }
        
        # Allocate tokens based on priority
        allocations = {
            'characters': 0.15,  # 15% for character info
            'events': 0.10,     # 10% for timeline
            'chapters': 0.35,   # 35% for chapter content
            'knowledge': 0.25,  # 25% for knowledge base
            'history': 0.15    # 15% for conversation history
        }
        
        # Process each context type in priority order
        for ctx_type in priority_order:
            allocation = int(available_tokens * allocations.get(ctx_type, 0.1))
            
            if ctx_type == 'characters' and characters:
                context['characters'], tokens = self._fit_items(
                    characters, allocation, key='name'
                )
                token_usage['characters'] = tokens
                
            elif ctx_type == 'events' and events:
                context['events'], tokens = self._fit_items(
                    events, allocation, key='title'
                )
                token_usage['events'] = tokens
                
            elif ctx_type == 'chapters' and chapters:
                context['chapters'], tokens = self._fit_items(
                    chapters, allocation, key='title'
                )
                token_usage['chapters'] = tokens
                
            elif ctx_type == 'knowledge' and knowledge:
                context['knowledge'], tokens = self._fit_items(
                    knowledge, allocation, key='title'
                )
                token_usage['knowledge'] = tokens
                
            elif ctx_type == 'history' and conversation_history:
                context['history'], tokens = self._fit_messages(
                    conversation_history, allocation
                )
                token_usage['history'] = tokens
        
        token_usage['total'] = sum(v for k, v in token_usage.items() if k != 'total')
        
        return {
            'context': context,
            'token_usage': token_usage,
            'available_for_response': self.max_tokens - token_usage['total'] - self.reserved_tokens
        }
    
    def _fit_items(self, items: List[Dict], max_tokens: int, 
                   key: str = 'content') -> tuple:
        """Fit items within token limit."""
        fitted = []
        total_tokens = 0
        
        for item in items:
            item_str = str(item)
            item_tokens = self.count_tokens(item_str)
            
            if total_tokens + item_tokens <= max_tokens:
                fitted.append(item)
                total_tokens += item_tokens
            else:
                # Try to fit a truncated version
                if 'content' in item:
                    remaining = max_tokens - total_tokens
                    if remaining > 100:  # Only if meaningful space left
                        truncated = self._truncate_to_tokens(
                            item['content'], remaining - 50
                        )
                        fitted.append({**item, 'content': truncated + '...'})
                        total_tokens += remaining
                break
        
        return fitted, total_tokens
    
    def _fit_messages(self, messages: List[Dict[str, str]], 
                      max_tokens: int) -> tuple:
        """Fit conversation messages within token limit, prioritizing recent."""
        # Start from most recent
        fitted = []
        total_tokens = 0
        
        for msg in reversed(messages):
            msg_str = f"{msg['role']}: {msg['content']}"
            msg_tokens = self.count_tokens(msg_str)
            
            if total_tokens + msg_tokens <= max_tokens:
                fitted.insert(0, msg)
                total_tokens += msg_tokens
            else:
                break
        
        return fitted, total_tokens
    
    def _truncate_to_tokens(self, text: str, max_tokens: int) -> str:
        """Truncate text to fit within token limit."""
        tokens = self.encoding.encode(text)
        if len(tokens) <= max_tokens:
            return text
        return self.encoding.decode(tokens[:max_tokens])


def get_document_processor() -> DocumentProcessor:
    """Get document processor instance."""
    return DocumentProcessor()


def get_long_context_manager(max_tokens: int = 32000) -> LongContextManager:
    """Get long context manager instance."""
    return LongContextManager(max_tokens)
