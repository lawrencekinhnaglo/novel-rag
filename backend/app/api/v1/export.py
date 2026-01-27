"""Export API - Export stories, chapters, and worldbuilding to various formats."""
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import text
from typing import Optional, List
from pydantic import BaseModel
from app.database.postgres import get_db
import io
import json
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

router = APIRouter()


class ExportRequest(BaseModel):
    """Export request parameters."""
    format: str = "docx"  # docx, pdf, markdown, json
    include_metadata: bool = True
    include_worldbuilding: bool = True
    include_chapters: bool = True
    include_characters: bool = True
    include_timeline: bool = False
    language: str = "zh-TW"


def generate_docx(title: str, content: dict) -> bytes:
    """Generate a DOCX file from content."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    if content.get("metadata"):
        meta = content["metadata"]
        doc.add_paragraph(f"Generated: {meta.get('generated_at', '')}")
        doc.add_paragraph(f"Language: {meta.get('language', 'zh-TW')}")
        doc.add_paragraph()
    
    # Series Info
    if content.get("series"):
        series = content["series"]
        doc.add_heading("Series Information", 1)
        doc.add_paragraph(f"Title: {series.get('title', 'Unknown')}")
        if series.get("premise"):
            doc.add_paragraph(f"Premise: {series['premise']}")
        if series.get("themes"):
            themes = series["themes"]
            if isinstance(themes, list):
                doc.add_paragraph(f"Themes: {', '.join(themes)}")
        doc.add_paragraph()
    
    # Worldbuilding
    if content.get("worldbuilding"):
        doc.add_heading("Worldbuilding", 1)
        for item in content["worldbuilding"]:
            doc.add_heading(item.get("title", "Untitled"), 2)
            doc.add_paragraph(item.get("content", ""))
            if item.get("category"):
                doc.add_paragraph(f"Category: {item['category']}", style="Intense Quote")
            doc.add_paragraph()
    
    # Characters
    if content.get("characters"):
        doc.add_heading("Characters", 1)
        for char in content["characters"]:
            doc.add_heading(char.get("name", "Unknown"), 2)
            if char.get("aliases"):
                doc.add_paragraph(f"Aliases: {', '.join(char['aliases'])}")
            if char.get("description"):
                doc.add_paragraph(char["description"])
            if char.get("personality"):
                doc.add_paragraph(f"Personality: {char['personality']}")
            if char.get("background"):
                doc.add_paragraph(f"Background: {char['background']}")
            doc.add_paragraph()
    
    # Chapters
    if content.get("chapters"):
        doc.add_heading("Chapters", 1)
        for chapter in content["chapters"]:
            ch_num = chapter.get("chapter_number", "?")
            ch_title = chapter.get("title", "Untitled")
            doc.add_heading(f"Chapter {ch_num}: {ch_title}", 2)
            doc.add_paragraph(chapter.get("content", ""))
            doc.add_paragraph()
    
    # World Rules
    if content.get("world_rules"):
        doc.add_heading("World Rules", 1)
        for rule in content["world_rules"]:
            doc.add_heading(rule.get("rule_name", "Rule"), 3)
            doc.add_paragraph(f"Category: {rule.get('rule_category', 'general')}")
            doc.add_paragraph(rule.get("rule_description", ""))
            doc.add_paragraph()
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_markdown(title: str, content: dict) -> str:
    """Generate Markdown from content."""
    lines = []
    
    # Title
    lines.append(f"# {title}")
    lines.append("")
    
    # Metadata
    if content.get("metadata"):
        meta = content["metadata"]
        lines.append(f"*Generated: {meta.get('generated_at', '')}*")
        lines.append("")
    
    # Series Info
    if content.get("series"):
        series = content["series"]
        lines.append("## Series Information")
        lines.append(f"**Title:** {series.get('title', 'Unknown')}")
        if series.get("premise"):
            lines.append(f"**Premise:** {series['premise']}")
        if series.get("themes"):
            themes = series["themes"]
            if isinstance(themes, list):
                lines.append(f"**Themes:** {', '.join(themes)}")
        lines.append("")
    
    # Worldbuilding
    if content.get("worldbuilding"):
        lines.append("## Worldbuilding")
        lines.append("")
        for item in content["worldbuilding"]:
            lines.append(f"### {item.get('title', 'Untitled')}")
            if item.get("category"):
                lines.append(f"*Category: {item['category']}*")
            lines.append("")
            lines.append(item.get("content", ""))
            lines.append("")
    
    # Characters
    if content.get("characters"):
        lines.append("## Characters")
        lines.append("")
        for char in content["characters"]:
            lines.append(f"### {char.get('name', 'Unknown')}")
            if char.get("aliases"):
                lines.append(f"**Aliases:** {', '.join(char['aliases'])}")
            if char.get("description"):
                lines.append(f"\n{char['description']}")
            if char.get("personality"):
                lines.append(f"\n**Personality:** {char['personality']}")
            if char.get("background"):
                lines.append(f"\n**Background:** {char['background']}")
            lines.append("")
    
    # Chapters
    if content.get("chapters"):
        lines.append("## Chapters")
        lines.append("")
        for chapter in content["chapters"]:
            ch_num = chapter.get("chapter_number", "?")
            ch_title = chapter.get("title", "Untitled")
            lines.append(f"### Chapter {ch_num}: {ch_title}")
            lines.append("")
            lines.append(chapter.get("content", ""))
            lines.append("")
            lines.append("---")
            lines.append("")
    
    # World Rules
    if content.get("world_rules"):
        lines.append("## World Rules")
        lines.append("")
        for rule in content["world_rules"]:
            lines.append(f"### {rule.get('rule_name', 'Rule')}")
            lines.append(f"*Category: {rule.get('rule_category', 'general')}*")
            lines.append("")
            lines.append(rule.get("rule_description", ""))
            lines.append("")
    
    return "\n".join(lines)


@router.post("/series/{series_id}")
async def export_series(
    series_id: int,
    request: ExportRequest,
    db: AsyncSession = Depends(get_db)
):
    """Export a complete series with all its content."""
    
    content = {
        "metadata": {
            "generated_at": datetime.now().isoformat(),
            "language": request.language,
            "export_format": request.format
        }
    }
    
    # Get series info
    series_result = await db.execute(
        text("SELECT id, title, premise, themes, language FROM series WHERE id = :id"),
        {"id": series_id}
    )
    series = series_result.fetchone()
    
    if not series:
        raise HTTPException(404, "Series not found")
    
    content["series"] = {
        "id": series.id,
        "title": series.title,
        "premise": series.premise,
        "themes": series.themes,
        "language": series.language
    }
    
    title = series.title or "Exported Series"
    
    # Get worldbuilding
    if request.include_worldbuilding:
        wb_result = await db.execute(
            text("""
                SELECT id, title, content, category, source_type
                FROM knowledge_base
                WHERE :series_tag = ANY(tags)
                ORDER BY created_at DESC
                LIMIT 100
            """),
            {"series_tag": f"series:{series_id}"}
        )
        content["worldbuilding"] = [
            {
                "id": row.id,
                "title": row.title,
                "content": row.content,
                "category": row.category or row.source_type
            }
            for row in wb_result.fetchall()
        ]
    
    # Get characters
    if request.include_characters:
        chars_result = await db.execute(
            text("""
                SELECT name, aliases, description, personality, background, goals, role
                FROM character_profiles
                WHERE series_id = :series_id
                ORDER BY first_appearance_book ASC NULLS LAST, created_at ASC
            """),
            {"series_id": series_id}
        )
        content["characters"] = [
            {
                "name": row.name,
                "aliases": row.aliases or [],
                "description": row.description,
                "personality": row.personality,
                "background": row.background,
                "goals": row.goals,
                "role": row.role
            }
            for row in chars_result.fetchall()
        ]
    
    # Get chapters
    if request.include_chapters:
        chapters_result = await db.execute(
            text("""
                SELECT c.chapter_number, c.title, c.content, c.synopsis, b.title as book_title, b.book_number
                FROM chapters c
                JOIN books b ON c.book_id = b.id
                WHERE b.series_id = :series_id
                ORDER BY b.book_number ASC, c.chapter_number ASC
            """),
            {"series_id": series_id}
        )
        content["chapters"] = [
            {
                "book_number": row.book_number,
                "book_title": row.book_title,
                "chapter_number": row.chapter_number,
                "title": row.title,
                "content": row.content,
                "synopsis": row.synopsis
            }
            for row in chapters_result.fetchall()
        ]
    
    # Get world rules
    rules_result = await db.execute(
        text("""
            SELECT rule_name, rule_category, rule_description
            FROM world_rules
            WHERE series_id = :series_id
        """),
        {"series_id": series_id}
    )
    content["world_rules"] = [
        {
            "rule_name": row.rule_name,
            "rule_category": row.rule_category,
            "rule_description": row.rule_description
        }
        for row in rules_result.fetchall()
    ]
    
    # Generate output based on format
    if request.format == "docx":
        output = generate_docx(title, content)
        filename = f"{title.replace(' ', '_')}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif request.format == "markdown":
        output = generate_markdown(title, content).encode("utf-8")
        filename = f"{title.replace(' ', '_')}.md"
        media_type = "text/markdown"
    elif request.format == "json":
        output = json.dumps(content, ensure_ascii=False, indent=2).encode("utf-8")
        filename = f"{title.replace(' ', '_')}.json"
        media_type = "application/json"
    else:
        raise HTTPException(400, f"Unsupported format: {request.format}")
    
    return StreamingResponse(
        io.BytesIO(output),
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename}\""
        }
    )


@router.post("/knowledge")
async def export_knowledge(
    request: ExportRequest,
    category: Optional[str] = None,
    series_id: Optional[int] = None,
    db: AsyncSession = Depends(get_db)
):
    """Export knowledge base entries."""
    
    # Build query
    query = "SELECT id, title, content, category, source_type, tags, created_at FROM knowledge_base WHERE 1=1"
    params = {}
    
    if category:
        query += " AND (category = :category OR source_type = :category)"
        params["category"] = category
    
    if series_id:
        query += " AND :series_tag = ANY(tags)"
        params["series_tag"] = f"series:{series_id}"
    
    query += " ORDER BY created_at DESC LIMIT 500"
    
    result = await db.execute(text(query), params)
    rows = result.fetchall()
    
    content = {
        "metadata": {
            "generated_at": datetime.now().isoformat(),
            "total_items": len(rows),
            "language": request.language
        },
        "worldbuilding": [
            {
                "id": row.id,
                "title": row.title,
                "content": row.content,
                "category": row.category or row.source_type,
                "tags": row.tags,
                "created_at": row.created_at.isoformat() if row.created_at else None
            }
            for row in rows
        ]
    }
    
    title = "Knowledge_Export"
    
    if request.format == "docx":
        output = generate_docx(title, content)
        filename = f"{title}_{datetime.now().strftime('%Y%m%d')}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif request.format == "markdown":
        output = generate_markdown(title, content).encode("utf-8")
        filename = f"{title}_{datetime.now().strftime('%Y%m%d')}.md"
        media_type = "text/markdown"
    elif request.format == "json":
        output = json.dumps(content, ensure_ascii=False, indent=2).encode("utf-8")
        filename = f"{title}_{datetime.now().strftime('%Y%m%d')}.json"
        media_type = "application/json"
    else:
        raise HTTPException(400, f"Unsupported format: {request.format}")
    
    return StreamingResponse(
        io.BytesIO(output),
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename}\""
        }
    )


@router.get("/formats")
async def list_formats():
    """List available export formats."""
    return {
        "formats": [
            {
                "id": "docx",
                "name": "Microsoft Word",
                "extension": ".docx",
                "description": "Editable document format"
            },
            {
                "id": "markdown",
                "name": "Markdown",
                "extension": ".md",
                "description": "Plain text with formatting"
            },
            {
                "id": "json",
                "name": "JSON",
                "extension": ".json",
                "description": "Structured data format"
            }
        ]
    }
